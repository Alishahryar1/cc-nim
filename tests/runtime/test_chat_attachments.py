import asyncio
import threading
import uuid
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from free_claude_code.application.chat import (
    ChatAttachment,
    ChatAttachmentKind,
    ChatDocumentAttachment,
    ChatImageAttachment,
    ChatPayloadTooLargeError,
    ChatUnsupportedAttachmentError,
    ChatValidationError,
)
from free_claude_code.runtime import chat_attachments
from free_claude_code.runtime.chat_attachments import LocalChatAttachmentFiles


def _id() -> str:
    return str(uuid.uuid4())


def _attachment(
    *,
    session_id: str,
    attachment_id: str,
    filename: str,
    kind: ChatAttachmentKind,
    media_type: str,
    byte_size: int,
    extracted_characters: int | None,
) -> ChatAttachment:
    return ChatAttachment(
        id=attachment_id,
        session_id=session_id,
        turn_id=None,
        position=0,
        filename=filename,
        kind=kind,
        media_type=media_type,
        byte_size=byte_size,
        extracted_characters=extracted_characters,
        created_at=1,
    )


def _pdf_bytes(*, pages: int = 1, encrypted: bool = False) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    for index in range(pages):
        page = writer.add_blank_page(width=200, height=200)
        if index:
            continue
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 20 100 Td (Hello PDF) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(stream)
    if encrypted:
        writer.encrypt("secret")
    writer.write(output)
    return output.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Project brief")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Ada"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _image_bytes(format_name: str) -> bytes:
    output = BytesIO()
    Image.new("RGB", (2, 2), color="orange").save(output, format=format_name)
    return output.getvalue()


@pytest.mark.asyncio
async def test_text_upload_is_verified_materialized_and_deleted(tmp_path: Path):
    files = LocalChatAttachmentFiles(tmp_path)
    session_id = _id()
    attachment_id = _id()
    await files.start(())

    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="notes.md",
        declared_media_type="text/markdown; charset=utf-8",
        source=BytesIO(b"heading\r\nbody"),
    )
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="notes.md",
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )

    assert info.kind is ChatAttachmentKind.MARKDOWN
    assert info.extracted_characters == len("heading\nbody")
    material = (await files.materialize((attachment,)))[0]
    assert isinstance(material, ChatDocumentAttachment)
    assert material.text == "heading\nbody"
    assert (await files.content(attachment)).data == b"heading\r\nbody"
    assert await files.available_ids((attachment,)) == frozenset({attachment_id})

    await files.delete_attachment(attachment)
    assert await files.available_ids((attachment,)) == frozenset()


@pytest.mark.asyncio
async def test_docx_upload_extracts_paragraphs_and_tables(tmp_path: Path):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    session_id = _id()
    attachment_id = _id()

    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="brief.docx",
        declared_media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        source=BytesIO(_docx_bytes()),
    )

    assert info.kind is ChatAttachmentKind.DOCX
    assert info.extracted_characters is not None
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="brief.docx",
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )
    material = (await files.materialize((attachment,)))[0]
    assert isinstance(material, ChatDocumentAttachment)
    assert material.text == "Project brief\nOwner\tAda"


@pytest.mark.parametrize(
    ("filename", "media_type", "format_name"),
    [
        ("photo.jpg", "image/jpeg", "JPEG"),
        ("photo.png", "image/png", "PNG"),
        ("photo.gif", "image/gif", "GIF"),
        ("photo.webp", "image/webp", "WEBP"),
    ],
)
@pytest.mark.asyncio
async def test_supported_image_signatures_are_preserved_as_original_bytes(
    tmp_path: Path,
    filename: str,
    media_type: str,
    format_name: str,
):
    data = _image_bytes(format_name)
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    session_id = _id()
    attachment_id = _id()

    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename=filename,
        declared_media_type=media_type,
        source=BytesIO(data),
    )
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename=filename,
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )

    material = (await files.materialize((attachment,)))[0]
    assert isinstance(material, ChatImageAttachment)
    assert material.data == data


@pytest.mark.parametrize(
    ("filename", "media_type", "data"),
    [
        ("broken.jpg", "image/jpeg", b"\xff\xd8\xfffixture"),
        ("broken.png", "image/png", b"\x89PNG\r\n\x1a\nfixture"),
        ("broken.gif", "image/gif", b"GIF89afixture"),
        ("broken.webp", "image/webp", b"RIFF\x08\x00\x00\x00WEBPfixture"),
    ],
)
@pytest.mark.asyncio
async def test_image_upload_rejects_corrupt_content_after_a_valid_signature(
    tmp_path: Path,
    filename: str,
    media_type: str,
    data: bytes,
):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())

    with pytest.raises(ChatUnsupportedAttachmentError, match=r"image.*malformed"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename=filename,
            declared_media_type=media_type,
            source=BytesIO(data),
        )


@pytest.mark.asyncio
async def test_image_upload_rejects_decompression_bomb_dimensions(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr(Image, "MAX_IMAGE_PIXELS", 1)
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())

    with pytest.raises(ChatPayloadTooLargeError, match="dimensions are too large"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="oversized.png",
            declared_media_type="image/png",
            source=BytesIO(_image_bytes("PNG")),
        )


@pytest.mark.asyncio
async def test_pdf_upload_extracts_text_and_rejects_encryption_or_page_overflow(
    tmp_path: Path,
):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    session_id = _id()
    attachment_id = _id()

    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="brief.pdf",
        declared_media_type="application/pdf",
        source=BytesIO(_pdf_bytes()),
    )
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="brief.pdf",
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )
    material = (await files.materialize((attachment,)))[0]
    assert isinstance(material, ChatDocumentAttachment)
    assert material.text == "Hello PDF"

    with pytest.raises(ChatUnsupportedAttachmentError, match="Encrypted PDF"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="encrypted.pdf",
            declared_media_type="application/pdf",
            source=BytesIO(_pdf_bytes(encrypted=True)),
        )
    with pytest.raises(ChatPayloadTooLargeError, match="400 pages"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="long.pdf",
            declared_media_type="application/pdf",
            source=BytesIO(_pdf_bytes(pages=401)),
        )


@pytest.mark.asyncio
async def test_pdf_upload_stops_extracting_as_soon_as_text_limit_is_exceeded(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    extracted_pages: list[int] = []

    class FakePage:
        def __init__(self, index: int) -> None:
            self._index = index

        def extract_text(self) -> str:
            extracted_pages.append(self._index)
            return "abc"

    class FakeReader:
        is_encrypted = False
        pages = tuple(FakePage(index) for index in range(3))

    def fake_reader(_path: Path) -> FakeReader:
        return FakeReader()

    monkeypatch.setattr(chat_attachments, "PdfReader", fake_reader)
    monkeypatch.setattr(
        chat_attachments,
        "MAX_CHAT_ATTACHMENT_EXTRACTED_CHARACTERS",
        5,
    )
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())

    with pytest.raises(ChatPayloadTooLargeError, match="1,000,000"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="large.pdf",
            declared_media_type="application/pdf",
            source=BytesIO(b"%PDF-fixture"),
        )

    assert extracted_pages == [0, 1]


@pytest.mark.asyncio
async def test_upload_rejects_extension_or_declared_type_mismatch(tmp_path: Path):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())

    with pytest.raises(ChatUnsupportedAttachmentError, match="extension"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="notes.pdf",
            declared_media_type="application/pdf",
            source=BytesIO(_image_bytes("PNG")),
        )

    with pytest.raises(ChatUnsupportedAttachmentError, match="media type"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="notes.txt",
            declared_media_type="image/png",
            source=BytesIO(b"plain text"),
        )


@pytest.mark.asyncio
async def test_upload_caps_streamed_bytes_and_extracted_characters(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    monkeypatch.setattr(chat_attachments, "MAX_CHAT_ATTACHMENT_BYTES", 4)
    with pytest.raises(ChatPayloadTooLargeError, match="10 MiB"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="note.txt",
            declared_media_type="text/plain",
            source=BytesIO(b"12345"),
        )

    monkeypatch.setattr(chat_attachments, "MAX_CHAT_ATTACHMENT_BYTES", 100)
    monkeypatch.setattr(
        chat_attachments,
        "MAX_CHAT_ATTACHMENT_EXTRACTED_CHARACTERS",
        4,
    )
    with pytest.raises(ChatPayloadTooLargeError, match="1,000,000"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="note.txt",
            declared_media_type="text/plain",
            source=BytesIO(b"12345"),
        )


@pytest.mark.asyncio
async def test_docx_rejects_unsafe_members_and_declared_archive_expansion(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    unsafe = BytesIO()
    with zipfile.ZipFile(unsafe, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../outside", "bad")
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())

    with pytest.raises(ChatUnsupportedAttachmentError, match="unsafe path"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="unsafe.docx",
            declared_media_type="application/zip",
            source=BytesIO(unsafe.getvalue()),
        )

    monkeypatch.setattr(chat_attachments, "_MAX_DOCX_UNCOMPRESSED_BYTES", 1)
    with pytest.raises(ChatPayloadTooLargeError, match="50 MiB"):
        await files.store_upload(
            session_id=_id(),
            attachment_id=_id(),
            filename="large.docx",
            declared_media_type="application/zip",
            source=BytesIO(_docx_bytes()),
        )


@pytest.mark.asyncio
async def test_start_removes_only_unowned_generated_attachment_trees(tmp_path: Path):
    files = LocalChatAttachmentFiles(tmp_path)
    session_id = _id()
    kept_id = _id()
    orphan_id = _id()
    await files.start(())
    for attachment_id in (kept_id, orphan_id):
        await files.store_upload(
            session_id=session_id,
            attachment_id=attachment_id,
            filename="note.txt",
            declared_media_type="text/plain",
            source=BytesIO(b"hello"),
        )

    await files.start(((session_id, kept_id),))

    attachment_root = tmp_path / "attachments" / session_id
    assert (attachment_root / kept_id).is_dir()
    assert not (attachment_root / orphan_id).exists()


@pytest.mark.asyncio
async def test_missing_owned_file_is_reported_unavailable_not_recreated(
    tmp_path: Path,
):
    files = LocalChatAttachmentFiles(tmp_path)
    session_id = _id()
    attachment_id = _id()
    await files.start(())
    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="note.txt",
        declared_media_type="text/plain",
        source=BytesIO(b"hello"),
    )
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="note.txt",
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )
    original = tmp_path / "attachments" / session_id / attachment_id / "original"
    original.unlink()

    await files.start(((session_id, attachment_id),))

    assert original.parent.is_dir()
    assert await files.available_ids((attachment,)) == frozenset()
    with pytest.raises(ChatValidationError, match="unavailable"):
        await files.content(attachment)


@pytest.mark.asyncio
async def test_cancelled_upload_waits_for_storage_and_removes_unowned_files(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    session_id = _id()
    attachment_id = _id()
    entered = threading.Event()
    release = threading.Event()
    original_store = files._store_upload_sync

    def blocked_store(**kwargs):
        entered.set()
        release.wait()
        return original_store(**kwargs)

    monkeypatch.setattr(files, "_store_upload_sync", blocked_store)
    upload = asyncio.create_task(
        files.store_upload(
            session_id=session_id,
            attachment_id=attachment_id,
            filename="note.txt",
            declared_media_type="text/plain",
            source=BytesIO(b"hello"),
        )
    )
    assert await asyncio.wait_for(asyncio.to_thread(entered.wait), timeout=1)

    upload.cancel()
    await asyncio.sleep(0.05)
    assert not upload.done()
    release.set()
    with pytest.raises(asyncio.CancelledError):
        await upload

    assert not (tmp_path / "tmp" / attachment_id).exists()
    assert not (tmp_path / "attachments" / session_id / attachment_id).exists()


@pytest.mark.asyncio
async def test_cancelled_materialization_waits_for_worker_to_settle(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    files = LocalChatAttachmentFiles(tmp_path)
    await files.start(())
    session_id = _id()
    attachment_id = _id()
    info = await files.store_upload(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="note.txt",
        declared_media_type="text/plain",
        source=BytesIO(b"hello"),
    )
    attachment = _attachment(
        session_id=session_id,
        attachment_id=attachment_id,
        filename="note.txt",
        kind=info.kind,
        media_type=info.media_type,
        byte_size=info.byte_size,
        extracted_characters=info.extracted_characters,
    )
    entered = threading.Event()
    release = threading.Event()
    original_materialize = files._materialize_sync

    def blocked_materialize(attachments):
        entered.set()
        release.wait()
        return original_materialize(attachments)

    monkeypatch.setattr(files, "_materialize_sync", blocked_materialize)
    materialize = asyncio.create_task(files.materialize((attachment,)))
    assert await asyncio.wait_for(asyncio.to_thread(entered.wait), timeout=1)

    materialize.cancel()
    await asyncio.sleep(0.05)
    assert not materialize.done()
    release.set()
    with pytest.raises(asyncio.CancelledError):
        await materialize
